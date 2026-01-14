from __future__ import annotations

from io import BytesIO
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Inches

def _sync_primary_from_contact_items(cv: dict) -> dict:
    if not isinstance(cv, dict):
        return cv

    contacts = cv.get("contact_items", [])
    if not isinstance(contacts, list):
        return cv

    type_to_field = {
        "email": "email",
        "phone": "telefon",
        "location": "adresa",
        "linkedin": "linkedin",
        "github": "github",
        "website": "website",
    }

    for t, field in type_to_field.items():
        if str(cv.get(field, "")).strip():
            continue
        for it in contacts:
            if not isinstance(it, dict):
                continue
            if it.get("type") != t:
                continue
            val = str(it.get("value", "")).strip()
            if val:
                cv[field] = val
                break

    # backfill names
    if not str(cv.get("nume_prenume", "")).strip() and str(cv.get("full_name", "")).strip():
        cv["nume_prenume"] = cv["full_name"]
    if not str(cv.get("full_name", "")).strip() and str(cv.get("nume_prenume", "")).strip():
        cv["full_name"] = cv["nume_prenume"]

    return cv


def _get_photo_bytes(cv: dict) -> Optional[bytes]:
    b = cv.get("photo")
    if isinstance(b, (bytes, bytearray)) and len(b) > 0:
        return bytes(b)
    return None


def _split_lines(text: str) -> List[str]:
    if not text:
        return []
    return [x.strip() for x in str(text).splitlines() if x.strip()]


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text or "", level=level)


def _add_bullets(doc: Document, lines: List[str]):
    for line in lines:
        t = (line or "").strip()
        if not t:
            continue
        t = t.lstrip("-•* ").strip()
        doc.add_paragraph(t, style="List Bullet")


def _skills_map_from_ats(cv: dict) -> Dict[str, List[str]]:
    out: Dict[str, List[str]] = {}
    skills = cv.get("ats_skills", [])
    if isinstance(skills, list):
        for sec in skills:
            if not isinstance(sec, dict):
                continue
            cat = (sec.get("category") or "").strip()
            items = sec.get("items", [])
            if cat and isinstance(items, list):
                out[cat.lower()] = [str(x).strip() for x in items if str(x).strip()]
    return out


def _build_technical_skills_lines(cv: dict) -> List[str]:
    m = _skills_map_from_ats(cv)
    if not m:
        m = {}
        if cv.get("modern_tools"):
            m["tools"] = _split_lines(cv.get("modern_tools", ""))
        if cv.get("modern_certs"):
            m["certifications"] = _split_lines(cv.get("modern_certs", ""))
        if cv.get("modern_keywords_extra"):
            m["keywords"] = _split_lines(cv.get("modern_keywords_extra", ""))

    def uniq(items: List[str]) -> List[str]:
        seen = set()
        out = []
        for x in items:
            k = x.strip()
            if not k:
                continue
            lk = k.lower()
            if lk in seen:
                continue
            seen.add(lk)
            out.append(k)
        return out

    cloud_identity = uniq(m.get("cloud", []) + m.get("identity", []) + m.get("cloud & identity", []))
    security = uniq(m.get("security", []) + m.get("keywords", []))
    networking = uniq(m.get("networking", []) + m.get("network", []))
    os_servers = uniq(m.get("windows/linux", []) + m.get("windows", []) + m.get("linux", []) + m.get("os & servers", []))
    scripting = uniq(m.get("scripting/automation", []) + m.get("automation", []) + m.get("scripting", []))
    tools = uniq(m.get("tools", []))
    virtualization = uniq(m.get("virtualization", []))

    sec_only, moved_to_net, moved_to_os = [], [], []
    for x in security:
        lx = x.lower()
        if any(k in lx for k in ["cisco", "vlan", "vpn", "firewall", "routing", "switch"]):
            moved_to_net.append(x)
        elif any(k in lx for k in ["windows", "linux", "active directory", "ad", "gpo", "server", "hyper-v", "vmware", "virtual"]):
            moved_to_os.append(x)
        else:
            sec_only.append(x)
    security = uniq(sec_only)
    networking = uniq(networking + moved_to_net)
    os_servers = uniq(os_servers + moved_to_os)

    for x in tools[:]:
        lx = x.lower()
        if any(k in lx for k in ["azure", "entra", "azure ad", "microsoft 365", "m365"]):
            cloud_identity.append(x)
            tools.remove(x)
    cloud_identity = uniq(cloud_identity)

    lines = []

    def add_group(title: str, items: List[str]):
        items = uniq(items)
        if items:
            lines.append(f"{title}: " + ", ".join(items))

    add_group("Cloud & Identity", cloud_identity)
    add_group("Security", security)
    add_group("Networking", networking)
    add_group("OS & Servers", os_servers)
    add_group("Scripting & Automation", scripting)
    add_group("Tools", tools)
    add_group("Virtualization", virtualization)

    certs = uniq(m.get("certifications", []))
    if certs:
        add_group("Certifications", certs)

    return lines


def generate_docx_modern(cv: dict, lang: str = "en") -> bytes:
    cv = _sync_primary_from_contact_items(cv)
    doc = Document()

    name = cv.get("nume_prenume") or cv.get("full_name") or "Full Name"
    p = doc.add_paragraph()
    r = p.add_run(str(name))
    r.bold = True

    profile_line = (cv.get("profile_line") or "").strip()
    if profile_line:
        doc.add_paragraph(profile_line)

    headline = cv.get("pozitie_vizata", "")
    if headline:
        doc.add_paragraph(str(headline))

    # contact line (short)
    extras = cv.get("personal_info_extra", [])
    city = ""
    availability = ""
    if isinstance(extras, list):
        for it in extras:
            if not isinstance(it, dict):
                continue
            lab = (it.get("label") or "").strip().lower()
            if lab in ("city", "oraș", "oras", "localitate"):
                city = it.get("value", "") or ""
            if lab in ("availability", "disponibilitate"):
                availability = it.get("value", "") or ""

    parts = []
    if city:
        parts.append(f"City: {city}")
    if availability:
        parts.append(f"Availability: {availability}")
    if parts:
        doc.add_paragraph(" | ".join(parts))

    links = []
    if cv.get("email"):
        links.append(f"Email: {cv.get('email')}")
    if cv.get("telefon"):
        links.append(f"Phone: {cv.get('telefon')}")
    if cv.get("linkedin"):
        links.append(f"LinkedIn: {cv.get('linkedin')}")
    if cv.get("github"):
        links.append(f"GitHub: {cv.get('github')}")
    if cv.get("website"):
        links.append(f"Website: {cv.get('website')}")
    if links:
        doc.add_paragraph(" | ".join(links))

    photo = _get_photo_bytes(cv)
    if cv.get("include_photo_modern") and photo:
        try:
            doc.add_picture(BytesIO(photo), width=Inches(1.2))
        except Exception:
            pass

    bullets = cv.get("rezumat_bullets", [])
    if isinstance(bullets, list) and bullets:
        _add_heading(doc, "SUMMARY", level=2)
        _add_bullets(doc, [str(x) for x in bullets])
    else:
        rez = cv.get("rezumat", "")
        if rez:
            _add_heading(doc, "SUMMARY", level=2)
            doc.add_paragraph(str(rez))

    tech_lines = cv.get("technical_skills_lines") or _build_technical_skills_lines(cv)
    if tech_lines:
        _add_heading(doc, "TECHNICAL SKILLS", level=2)
        _add_bullets(doc, tech_lines)

    exp = cv.get("experienta", [])
    if isinstance(exp, list) and exp:
        _add_heading(doc, "PROFESSIONAL EXPERIENCE", level=2)
        for e in exp:
            if not isinstance(e, dict):
                continue
            company = e.get("titlu") or e.get("angajator") or ""
            period = e.get("perioada") or ""
            role = e.get("functie") or ""
            line = company
            if role:
                line = f"{role} — {company}" if company else role
            if period:
                line = f"{line} — {period}" if line else period

            p = doc.add_paragraph()
            rr = p.add_run(str(line))
            rr.bold = True

            loc = e.get("locatie", "")
            if loc:
                doc.add_paragraph(f"Location: {loc}")

            acts = _split_lines(e.get("activitati", ""))
            if acts:
                _add_bullets(doc, acts)

    edu = cv.get("educatie", [])
    if isinstance(edu, list) and edu:
        _add_heading(doc, "EDUCATION", level=2)
        for ed in edu:
            if not isinstance(ed, dict):
                continue
            per = ed.get("perioada", "")
            tit = ed.get("titlu", "")
            org = ed.get("organizatie", "")
            line = " — ".join([x for x in [per, tit] if x])
            doc.add_paragraph(line)
            if org:
                doc.add_paragraph(str(org))

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_docx_europass(cv: dict, lang: str = "en") -> bytes:
    cv = _sync_primary_from_contact_items(cv)
    doc = Document()

    name = cv.get("nume_prenume") or cv.get("full_name") or "Full Name"
    p = doc.add_paragraph()
    r = p.add_run(str(name))
    r.bold = True

    doc.add_paragraph("CURRICULUM VITAE (Europass)")

    profile_line = (cv.get("profile_line") or "").strip()
    if profile_line:
        doc.add_paragraph(profile_line)

    photo = _get_photo_bytes(cv)
    if photo:
        try:
            doc.add_picture(BytesIO(photo), width=Inches(1.3))
        except Exception:
            pass

    _add_heading(doc, "PERSONAL INFORMATION", level=2)

    extras = cv.get("personal_info_extra", [])
    city = ""
    availability = ""
    if isinstance(extras, list):
        for it in extras:
            if not isinstance(it, dict):
                continue
            lab = (it.get("label") or "").strip().lower()
            if lab in ("city", "oraș", "oras", "localitate"):
                city = it.get("value", "") or ""
            if lab in ("availability", "disponibilitate"):
                availability = it.get("value", "") or ""
    parts = []
    if city:
        parts.append(f"City: {city}")
    if availability:
        parts.append(f"Availability: {availability}")
    if parts:
        doc.add_paragraph(" | ".join(parts))

    lines = []
    if cv.get("email"):
        lines.append(f"Email: {cv.get('email')}")
    if cv.get("telefon"):
        lines.append(f"Phone: {cv.get('telefon')}")
    if cv.get("adresa"):
        lines.append(f"Location: {cv.get('adresa')}")
    if cv.get("linkedin"):
        lines.append(f"LinkedIn: {cv.get('linkedin')}")
    if cv.get("github"):
        lines.append(f"GitHub: {cv.get('github')}")
    if cv.get("website"):
        lines.append(f"Website: {cv.get('website')}")

    if isinstance(extras, list):
        for it in extras:
            if not isinstance(it, dict):
                continue
            lab = (it.get("label") or "").strip()
            val = (it.get("value") or "").strip()
            if lab and val and lab.lower() not in ("city", "oraș", "oras", "localitate", "availability", "disponibilitate"):
                lines.append(f"{lab}: {val}")

    _add_bullets(doc, lines)

    exp = cv.get("experienta", [])
    if isinstance(exp, list) and exp:
        _add_heading(doc, "WORK EXPERIENCE", level=2)
        for e in exp:
            if not isinstance(e, dict):
                continue
            role = e.get("functie") or ""
            company = e.get("angajator") or e.get("titlu") or ""
            period = e.get("perioada") or ""
            title = " — ".join([x for x in [role, period] if x])

            p = doc.add_paragraph()
            rr = p.add_run(title)
            rr.bold = True
            if company:
                doc.add_paragraph(str(company))
            loc = e.get("locatie") or ""
            if loc:
                doc.add_paragraph(f"Location: {loc}")

            acts = _split_lines(e.get("activitati", ""))
            _add_bullets(doc, acts)

    edu = cv.get("educatie", [])
    if isinstance(edu, list) and edu:
        _add_heading(doc, "EDUCATION AND TRAINING", level=2)
        for ed in edu:
            if not isinstance(ed, dict):
                continue
            per = ed.get("perioada", "")
            tit = ed.get("titlu", "")
            org = ed.get("organizatie", "")
            line = " — ".join([x for x in [per, tit] if x])
            doc.add_paragraph(line)
            if org:
                doc.add_paragraph(str(org))

    _add_heading(doc, "LANGUAGE SKILLS", level=2)
    mother = cv.get("limba_materna", "")
    if mother:
        doc.add_paragraph(f"Mother tongue: {mother}")

    langs = cv.get("limbi_straine", [])
    if isinstance(langs, list) and langs:
        lines = []
        for l in langs:
            if not isinstance(l, dict):
                continue
            nm = l.get("limba", "")
            lvl = l.get("nivel", "") or l.get("ascultare", "")
            if nm:
                lines.append(f"{nm}: {lvl}".strip())
        _add_bullets(doc, lines)

    _add_heading(doc, "PERSONAL SKILLS", level=2)
    secs = cv.get("aptitudini_sections", [])
    if isinstance(secs, list) and secs:
        for sec in secs:
            if not isinstance(sec, dict):
                continue
            cat = sec.get("category", "")
            items = sec.get("items", [])
            if cat:
                pp = doc.add_paragraph()
                rr = pp.add_run(cat)
                rr.bold = True
            if isinstance(items, list) and items:
                _add_bullets(doc, [str(x) for x in items])

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
