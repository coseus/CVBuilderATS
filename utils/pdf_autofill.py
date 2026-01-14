import re
from typing import Dict, List, Tuple, Optional

import pdfplumber

try:
    from docx import Document  # python-docx
except Exception:
    Document = None


# -----------------------------
# Helpers
# -----------------------------
def _clean(s: str) -> str:
    s = re.sub(r"\s+", " ", (s or "")).strip()
    return _dedupe_doubled_chars(s)


def _clean_lines(text: str) -> List[str]:
    lines = (text or "").replace("\r", "\n").split("\n")
    out = []
    for ln in lines:
        ln = ln.strip()
        if ln:
            out.append(ln)
    return out


def _find_first(pattern: str, text: str, flags=re.IGNORECASE) -> Optional[str]:
    m = re.search(pattern, text, flags)
    return _clean(m.group(1)) if m else None


def _normalize_phone(x: str) -> str:
    x = _clean(x)
    # keep + and digits/spaces
    x = re.sub(r"[^\d\+\s]", "", x)
    x = re.sub(r"\s+", " ", x).strip()
    return x


def _normalize_url(x: str) -> str:
    x = _clean(x)
    # remove trailing punctuation
    x = x.rstrip(".,;)")
    return x


def _looks_like_name(line: str) -> bool:
    # simple heuristic: 2-4 words starting with capitals
    if not line:
        return False
    if len(line) > 60:
        return False
    words = line.split()
    if len(words) < 2 or len(words) > 5:
        return False
    cap = 0
    for w in words:
        if w[:1].isupper():
            cap += 1
    return cap >= 2


def _extract_contact_urls(text: str) -> Dict[str, str]:
    """
    Extract linkedin/github/website from raw text using robust patterns.
    Always returns a dict.
    """
    txt = text or ""
    lower = txt.lower()

    def find_full(rx: str) -> str:
        m = re.search(rx, txt, flags=re.IGNORECASE)
        return _normalize_url(m.group(0)) if m else ""

    # LinkedIn / GitHub (full match, not group(1))
    linkedin = find_full(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s\)\],;]+")
    github = find_full(r"(?:https?://)?(?:www\.)?github\.com/[^\s\)\],;]+")

    # Normalize display (keep without scheme to be ATS-friendly if you prefer)
    if linkedin:
        linkedin = linkedin.replace("https://", "").replace("http://", "").rstrip("/")
    if github:
        github = github.replace("https://", "").replace("http://", "").rstrip("/")

    # Website:
    # Avoid email domains: negative lookbehind for '@'
    website = ""
    candidates = re.findall(
        r"(?<!@)\b(?:https?://)?(?:www\.)?[a-z0-9][a-z0-9\-]+\.[a-z]{2,}(?:/[^\s\)\],;]+)?\b",
        lower
    )
        # Blocklist for common CV platform footers / mail providers
    blocked_domains = {
        "ejobs.ro", "www.ejobs.ro",
        "contact@ejobs.ro",
        "linkedin.com", "github.com",  # we store these separately
        "yahoo.com", "gmail.com", "outlook.com", "hotmail.com", "live.com", "icloud.com"
    }

    if website:
        dom = website.lower().replace("https://", "").replace("http://", "")
        dom = dom.replace("www.", "")
        dom = dom.split("/")[0]
        if dom in blocked_domains or "ejobs.ro" in dom:
            website = ""

    if candidates:
        for c in candidates:
            if "linkedin.com" in c or "github.com" in c:
                continue
            website = _normalize_url(c)
            website = website.replace("https://", "").replace("http://", "").rstrip("/")
            break

    # Drop common mail providers that slip through
    mail_domains = {"yahoo.com", "gmail.com", "outlook.com", "hotmail.com", "live.com", "icloud.com"}
    if website:
        dom = website.lower().replace("www.", "").split("/")[0]
        if dom in mail_domains:
            website = ""

    return {
        "linkedin": linkedin or "",
        "github": github or "",
        "website": website or "",
    }




def _make_contact_items(email: str, phone: str, location: str, linkedin: str, github: str, website: str) -> List[Dict]:
    items = []
    def add(t, v, label=None):
        v = _clean(v)
        if v:
            items.append({"type": t, "value": v, "label": label or t})

    add("email", email, "email")
    add("phone", phone, "phone")
    add("location", location, "location")
    add("linkedin", linkedin, "linkedin")
    add("github", github, "github")
    add("website", website, "website")
    return items

def _dedupe_doubled_chars(s: str) -> str:
    """
    Fix PDFs that extract text with each character duplicated:
      'CCoossmmiinn' -> 'Cosmin'
    Applies per token; preserves whitespace.
    """
    s = (s or "").strip()
    if len(s) < 4:
        return s

    def dedupe_token(tok: str) -> str:
        t = tok.strip()
        if len(t) >= 4 and len(t) % 2 == 0:
            ok = True
            for i in range(0, len(t), 2):
                if t[i] != t[i + 1]:
                    ok = False
                    break
            if ok:
                return t[0::2]
        return tok

    parts = re.split(r"(\s+)", s)
    parts = [dedupe_token(p) if not p.isspace() else p for p in parts]
    return "".join(parts)

# -----------------------------
# Block extraction (multi-layout)
# -----------------------------
def _extract_blocks(text: str) -> Dict[str, any]:
    """
    Works for eJobs-like CVs and also more general layouts.
    Returns dict with keys:
      name, headline, profile_line, email, phone, location,
      about, education_lines (list), languages (list tuples), driving (list)
    """
    blocks: Dict[str, any] = {}

    # Email / phone / location (EN/RO)
    email = _find_first(r"\bEmail:\s*([^\s]+)", text) or _find_first(r"\bE-mail:\s*([^\s]+)", text)
    if not email:
        email = _find_first(r"\b([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})\b", text)

    phone = _find_first(r"(?:Tel|Telefon|Phone):\s*([+\d][\d\s\-\(\)]+)", text)
    phone = _normalize_phone(phone or "")

    location = (
        _find_first(r"(?:City|Oraș|Oras|Locație|Locatie):\s*([^\n]+)", text)
        or _find_first(r"(?:Address|Adresa|Adresă):\s*([^\n]+)", text)
        or ""
    )
    location = _clean(location)

    # Guess name/headline from top lines
    lines = _clean_lines(text)
    top = lines[:25]

    name = ""
    headline = ""
    profile_line = ""

    # name: first "name-like" line
    for ln in top:
        if _looks_like_name(ln):
            name = ln
            break

    # headline/profile line: usually next 1-3 lines after name, non-empty, not "City/Email"
    if name:
        try:
            idx = top.index(name)
        except ValueError:
            idx = 0
        candidates = top[idx+1:idx+6]
        for ln in candidates:
            lnl = ln.lower()
            if any(k in lnl for k in ["email:", "e-mail:", "phone", "telefon", "city:", "oraș", "oras", "location", "loca"]):
                continue
            if len(ln) < 3:
                continue
            # first meaningful line = headline
            if not headline:
                headline = ln
                continue
            # second meaningful line = profile_line (optional)
            if not profile_line and len(ln) <= 80:
                profile_line = ln
                break

    blocks["nume"] = name or ""
    blocks["headline"] = headline or ""
    blocks["profile_line"] = profile_line or ""
    blocks["email"] = email or ""
    blocks["telefon"] = phone or ""
    blocks["adresa"] = location or ""

    # URLs
    urls = _extract_contact_urls(text)
    blocks.update(urls if isinstance(urls, dict) else {})

    # About me / Despre mine / Summary
    about = ""
    about_patterns = [
        (r"About me\s*(.+?)\s*(Professional experience|Work experience|Experience|Employment history)", re.IGNORECASE | re.DOTALL),
        (r"Summary\s*(.+?)\s*(Professional experience|Work experience|Experience|Employment history)", re.IGNORECASE | re.DOTALL),
        (r"Despre mine\s*(.+?)\s*(Experiență profesională|Experienta profesionala)", re.IGNORECASE | re.DOTALL),
        (r"Rezumat\s*(.+?)\s*(Experiență profesională|Experienta profesionala)", re.IGNORECASE | re.DOTALL),
    ]
    for pat, flags in about_patterns:
        m = re.search(pat, text, flags)
        if m:
            about = _clean(m.group(1))
            break
    blocks["about"] = about or ""

    # Education lines: attempt to capture lines with year ranges
    edu_lines = []
    for ln in lines:
        if re.search(r"\b(19|20)\d{2}\s*-\s*(19|20)\d{2}\b", ln):
            # likely education or timeline entry; keep, filter later
            edu_lines.append(ln)
    blocks["education_lines"] = edu_lines

    # Foreign languages section
    langs: List[Tuple[str, str]] = []
    lang_section = None
    for pat in [
        r"Foreign languages\s*(.+?)\s*(?:Other sections|Other information|www\.ejobs\.ro|$)",
        r"Limbi străine\s*(.+?)\s*(?:Alte informații|Alte informatii|www\.ejobs\.ro|$)",
        r"Languages\s*(.+?)\s*(?:Other sections|Other information|$)",
        r"Limbi\s*(.+?)\s*(?:Alte informații|Alte informatii|$)",
    ]:
        m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
        if m:
            lang_section = m.group(1)
            break
    if lang_section:
        for raw in lang_section.splitlines():
            line = _clean(raw)
            if ":" in line:
                a, b = line.split(":", 1)
                a, b = _clean(a), _clean(b)
                if a and b and len(a) <= 35:
                    langs.append((a, b))
    blocks["languages"] = langs

    # Driving license
    dl = []
    for pat in [
        r"Driving license\s*(.+?)(?:www\.ejobs\.ro|$)",
        r"Permis de conducere\s*(.+?)(?:www\.ejobs\.ro|$)",
    ]:
        m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
        if m:
            seg = m.group(1)
            for c in re.findall(r"\bCategory\s*([A-Z])\b|\bCategoria\s*([A-Z])\b", seg, re.IGNORECASE):
                cat = (c[0] or c[1] or "").upper()
                if cat and cat not in dl:
                    dl.append(cat)
            break
    blocks["driving"] = dl

    return blocks


# -----------------------------
# Experience extraction (improved)
# -----------------------------
_MONTHS = r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Ian|Feb|Mar|Apr|Mai|Iun|Iul|Aug|Sep|Oct|Noi|Dec)"
_DATE_RANGE_RX = re.compile(
    rf"(?P<range>{_MONTHS}\s+\d{{4}}\s*-\s*(?:present|prezent|{_MONTHS}\s+\d{{4}}))",
    re.IGNORECASE,
)


def _extract_experience_items(text: str) -> List[Dict]:
    """
    More robust:
    - detects date ranges (month year - present/month year)
    - tries to parse next 1-2 lines as role/company
    - collects bullets until next range or footer
    """
    items: List[Dict] = []

    t = (text or "").replace("–", "-").replace("—", "-")
    lines = t.splitlines()

    # Build index of date-range lines
    indices = []
    for i, ln in enumerate(lines):
        if _DATE_RANGE_RX.search(ln):
            indices.append(i)

    if not indices:
        return items

    indices.append(len(lines))  # sentinel

    for idx_pos in range(len(indices) - 1):
        i = indices[idx_pos]
        j = indices[idx_pos + 1]

        rng = _clean(lines[i])
        # take next non-empty line(s)
        role_line = ""
        company_line = ""
        k = i + 1
        while k < j and not _clean(lines[k]):
            k += 1
        if k < j:
            role_line = _clean(lines[k])

        # sometimes role+company is in one line "Role - Company"
        role = role_line
        company = ""
        if " - " in role_line:
            role, company = role_line.split(" - ", 1)
            role, company = _clean(role), _clean(company)
        else:
            # try next line as company if short
            k2 = k + 1
            while k2 < j and not _clean(lines[k2]):
                k2 += 1
            if k2 < j:
                possible = _clean(lines[k2])
                # heuristics: company name tends to be shorter, no full stop
                if possible and len(possible) <= 60 and not possible.endswith("."):
                    company = possible

        # collect description lines between header and next entry
        desc_lines = lines[(k + 1):j]
        # stop at ejobs footer
        cut = []
        for ln in desc_lines:
            if "www.ejobs.ro" in ln.lower():
                break
            cut.append(ln)
        desc_lines = cut

        # clean acquired skills blocks
        desc_text = "\n".join(desc_lines)
        desc_text = re.sub(r"Acquired skills and competencies:.*", "", desc_text, flags=re.IGNORECASE | re.DOTALL)
        desc_text = re.sub(r"Abilități și competențe dobândite:.*", "", desc_text, flags=re.IGNORECASE | re.DOTALL)

        bullets: List[str] = []
        for ln in desc_text.splitlines():
            s = _clean(ln)
            if not s:
                continue
            if s.startswith("-") or s.startswith("•"):
                bullets.append(_clean(s.lstrip("-• ").strip()))

        if not bullets:
            # fallback: split longer sentences, keep strong ones
            s = _clean(desc_text)
            parts = re.split(r"\.\s+", s)
            bullets = [p.strip().rstrip(".") for p in parts if len(p.strip()) >= 35][:7]

        items.append({
            "titlu": company or role,
            "perioada": rng,
            "functie": role,
            "angajator": company,
            "locatie": "",
            "activitati": "\n".join([f"- {b}" for b in bullets if b]),
            "sector": "",
            "tehnologii": "",
            "link": "",
        })

    return items


# -----------------------------
# Education extraction (improved)
# -----------------------------
def _extract_education(blocks: Dict[str, any], full_text: str) -> List[Dict]:
    educatie: List[Dict] = []

    # 1) try "Education" section snippet (EN/RO)
    sec = None
    for pat in [
        r"(Education|Education and training)\s*(.+?)(?:Foreign languages|Languages|Other|www\.ejobs\.ro|$)",
        r"(Educație|Educatie)\s*(.+?)(?:Limbi străine|Limbi straine|Limbi|Alte|www\.ejobs\.ro|$)",
    ]:
        m = re.search(pat, full_text, re.IGNORECASE | re.DOTALL)
        if m:
            sec = m.group(2)
            break

    candidates = []
    if sec:
        candidates.extend(_clean_lines(sec))

    # 2) fallback: year-range lines we saw
    candidates.extend(blocks.get("education_lines", []) or [])

    # filter likely edu lines (avoid experience date ranges with months)
    for ln in candidates:
        if re.search(rf"\b{_MONTHS}\b", ln, re.IGNORECASE):
            continue
        if not re.search(r"\b(19|20)\d{2}\s*-\s*(19|20)\d{2}\b", ln):
            continue

        m = re.match(r"((?:19|20)\d{2}\s*-\s*(?:19|20)\d{2})\s+(.*)", ln)
        perio = _clean(m.group(1)) if m else ""
        rest = _clean(m.group(2)) if m else _clean(ln)

        # split " - " for title vs org if present
        titlu = rest
        org = ""
        if " - " in rest:
            left, right = rest.split(" - ", 1)
            titlu = _clean(left)
            org = _clean(right)

        # avoid duplicates
        key = (perio, titlu, org)
        if any((e.get("perioada"), e.get("titlu"), e.get("organizatie")) == key for e in educatie):
            continue

        educatie.append({
            "perioada": perio,
            "titlu": titlu,
            "organizatie": org,
            "locatie": "",
            "descriere": "",
        })

    return educatie[:5]


# -----------------------------
# Summary bullets
# -----------------------------
def _summary_to_bullets(about: str) -> List[str]:
    about = _clean(about)
    if not about:
        return []
    parts = re.split(r"\.\s+", about)
    bullets = []
    for p in parts:
        p = _clean(p).rstrip(".")
        if len(p) >= 25:
            bullets.append(p)
    return bullets[:5]


# -----------------------------
# Text loaders (PDF / DOCX)
# -----------------------------
def _read_pdf_text(pdf_path: str) -> str:
    pages_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for p in pdf.pages:
            pages_text.append(p.extract_text() or "")
    return "\n".join(pages_text)


def _read_docx_text(docx_path: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx not available. Add it to requirements / environment.")
    doc = Document(docx_path)
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # tables too (often used in Europass)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return "\n".join(parts)


# -----------------------------
# Public API
# -----------------------------
def text_to_cv(text: str, lang_hint: str = "en") -> Dict:
    blocks = _extract_blocks(text)
    exp = _extract_experience_items(text)
    educatie = _extract_education(blocks, text)

    # Summary bullets
    summary_bullets = _summary_to_bullets(blocks.get("about", ""))

    # Languages (simple level mapping)
    limbi = []
    for name, level in blocks.get("languages", []) or []:
        lvl = _clean(level)
        limbi.append({
            "limba": _clean(name),
            "nivel": lvl,
            "ascultare": lvl,
            "citire": lvl,
            "interactiune": lvl,
            "exprimare": lvl,
            "scriere": lvl,
        })

    driving = blocks.get("driving", []) or []

    # Contacts + contact_items
    email = blocks.get("email", "")
    phone = blocks.get("telefon", "")
    location = blocks.get("adresa", "")
    linkedin = blocks.get("linkedin", "")
    github = blocks.get("github", "")
    website = blocks.get("website", "")
    if email and website and website.lower().replace("www.", "") in email.lower():
        website = ""

    contact_items = _make_contact_items(email, phone, location, linkedin, github, website)

    # personal_info_extra: keep city/availability if present in text; we store location in adresa
    personal_extra = []
    # heuristic: if availability appears
    av = _find_first(r"(Availability|Disponibilitate)\s*:\s*([^\n]+)", text)
    if av:
        personal_extra.append({"label": "Availability", "value": av})

    cv = {
        "nume_prenume": blocks.get("nume", ""),
        "full_name": blocks.get("nume", ""),
        "profile_line": blocks.get("profile_line", ""),
        "pozitie_vizata": blocks.get("headline", ""),
        "email": email,
        "telefon": phone,
        "adresa": location,
        "linkedin": linkedin,
        "github": github,
        "website": website,
        "contact_items": contact_items,
        "rezumat_bullets": summary_bullets,
        "experienta": exp,
        "educatie": educatie,
        "limbi_straine": limbi,
        "permis_conducere": ", ".join(driving),
        "personal_info_extra": personal_extra,
        # Modern skills placeholders
        "modern_skills_headline": "",
        "modern_tools": "",
        "modern_certs": "",
        "modern_keywords_extra": "",
        # ATS profile default
        "ats_profile": "cyber_security",
    }
    return cv


def pdf_to_cv(pdf_path: str, lang_hint: str = "en") -> Dict:
    text = _read_pdf_text(pdf_path)
    return text_to_cv(text, lang_hint=lang_hint)


def docx_to_cv(docx_path: str, lang_hint: str = "en") -> Dict:
    text = _read_docx_text(docx_path)
    return text_to_cv(text, lang_hint=lang_hint)


def file_to_cv(path: str, lang_hint: str = "en") -> Dict:
    """
    Dispatch based on extension. Supports .pdf and .docx
    """
    p = (path or "").lower().strip()
    if p.endswith(".pdf"):
        return pdf_to_cv(path, lang_hint=lang_hint)
    if p.endswith(".docx"):
        return docx_to_cv(path, lang_hint=lang_hint)
    raise ValueError("Unsupported file type. Please upload PDF or DOCX.")
